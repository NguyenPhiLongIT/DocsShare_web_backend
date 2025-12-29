import re
from PyPDF2 import PdfReader
import fitz
from docx import Document

EXCLUDE_SECTIONS = [
    "lời cảm ơn", "lời mở đầu", "thông tin tác giả", "mục lục", 
    "tài liệu tham khảo", "phụ lục", "references", "acknowledgment",
    "thank you", "author", "contents"
]

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PdfReader(f)
            num_pages = len(reader.pages)

            # Nếu ít trang thì lấy hết
            if num_pages <= 3:
                start_page, end_page = 0, num_pages
            else:
                start_page, end_page = 1, num_pages - 1  # bỏ trang đầu/cuối

            for i in range(start_page, end_page):
                page = reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[PDF ERROR] {e}")
    return text.strip()

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        paragraphs = doc.paragraphs
        total = len(paragraphs)

        # Nếu ít paragraph thì lấy hết
        if total <= 10:
            start, end = 0, total
        else:
            start, end = 1, total - 1  # bỏ phần đầu/cuối thường là cảm ơn, tài liệu tham khảo

        for para in paragraphs[start:end]:
            if para.text.strip():
                text += para.text + "\n"
    except Exception as e:
        print(f"[DOCX ERROR] {e}")
    return text.strip()

def should_exclude(line):
    return any(kw in line.lower() for kw in EXCLUDE_SECTIONS)

def looks_like_heading(line):
    # Tiêu đề thường ngắn, viết hoa, có từ khóa như chương, phần
    return (
        len(line.strip()) < 80 and (
            re.match(r"^\s*(chương|phần)\s*\d+", line.lower()) or
            line.isupper()
        )
    )

def filter_excluded_blocks(text):
    lines = text.splitlines()
    output_lines = []
    skip = False

    for i in range(len(lines)):
        line = lines[i].strip()

        if should_exclude(line):
            skip = True
            continue

        if skip and looks_like_heading(line):
            skip = False  # Gặp heading mới thì dừng bỏ qua

        if not skip:
            output_lines.append(lines[i])

    return "\n".join(output_lines)

def extract_text(file_path):
    if file_path.lower().endswith(".pdf"):
        raw_text = extract_text_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        raw_text = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Only PDF and DOCX are supported.")
    
    return filter_excluded_blocks(raw_text)


def extract_first_sentences_pdf(file_path, n=5):
    full_text = []

    with fitz.open(file_path) as doc:
        for page in doc:
            text = page.get_text("text")
            if text:
                full_text.append(text.strip())

    full_text = " ".join(full_text)

    full_text = re.sub(r'\s+', ' ', full_text)

    sentences = re.split(r'(?<=[.!?])\s+', full_text)

    return " ".join(sentences[:n])

def extract_first_sentences_docx(file_path, n=5):
    doc = Document(file_path)

    full_text = " ".join(
        para.text.strip()
        for para in doc.paragraphs
        if para.text.strip()
    )

    sentences = re.split(r'(?<=[.!?])\s+', full_text)

    return " ".join(sentences[:n])

def extract_n_sentences(file_path):
    if file_path.lower().endswith(".pdf"):
        return extract_first_sentences_pdf(file_path, 5)
    elif file_path.lower().endswith(".docx"):
        return extract_first_sentences_docx(file_path, 5)
    else:
        raise ValueError("Unsupported file type")
